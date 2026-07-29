"""
extract_text.py
Day 5 Mission - Unstructured Data Layer

Pulls text from:
  1. A PDF policy document (pdfplumber)
  2. A Word claim form (python-docx)
  3. A scanned enrollment form image (pytesseract OCR)
  4. A public provider FAQ page (requests + BeautifulSoup)

Normalizes all extracted text (strip boilerplate, fix whitespace/encoding,
de-duplicate lines) and saves clean .txt files under raw_text/.

Usage:
    python extract_text.py
"""

import os
import re
import requests
import pdfplumber
import docx
import pytesseract
from PIL import Image
from bs4 import BeautifulSoup

SOURCE_DIR = "source_docs"
OUTPUT_DIR = "raw_text"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Normalization helpers
# ---------------------------------------------------------------------------

def fix_ocr_pipe_i(text: str) -> str:
    """
    OCR-specific fix: Tesseract frequently misreads a standalone capital "I"
    as a pipe character "|" on scanned forms. Safe to apply only to OCR
    output - other extractors (e.g. DOCX tables) use "|" intentionally as
    a column delimiter, so this must NOT run in the shared normalizer.
    """
    text = re.sub(r"(?<=\s)\|(?=\s)", "I", text)
    text = re.sub(r"^\|(?=\s)", "I", text, flags=re.MULTILINE)
    return text


def normalize_text(text: str, drop_lines_containing=None) -> str:
    """
    Clean up raw extracted text:
      - fix encoding artifacts / smart quotes
      - collapse repeated whitespace
      - strip lines that match known boilerplate (headers/footers/nav)
      - de-duplicate consecutive identical lines
      - trim leading/trailing blank lines
    """
    if not text:
        return ""

    # Fix common encoding artifacts
    replacements = {
        "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-", "\xa0": " ", "\u2022": "-",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)

    # pdfplumber sometimes can't map custom bullet glyphs and leaves the raw
    # PDF character code behind (e.g. "(cid:127)") - normalize to a plain dash
    text = re.sub(r"\(cid:\d+\)", "-", text)



    lines = text.split("\n")
    cleaned_lines = []
    prev_line = None

    drop_patterns = drop_lines_containing or []

    for raw_line in lines:
        line = re.sub(r"[ \t]+", " ", raw_line).strip()

        if not line:
            continue

        # Drop lines matching known boilerplate/nav patterns
        if any(pat.lower() in line.lower() for pat in drop_patterns):
            continue

        # Drop de-duplicated consecutive repeats (common in scraped nav menus)
        if line == prev_line:
            continue

        cleaned_lines.append(line)
        prev_line = line

    normalized = "\n".join(cleaned_lines)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def save_text(filename: str, text: str):
    path = os.path.join(OUTPUT_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"  -> saved {path} ({len(text)} chars)")


# ---------------------------------------------------------------------------
# 1. PDF extraction (pdfplumber)
# ---------------------------------------------------------------------------

def extract_pdf(path: str) -> str:
    print(f"[1/4] Extracting PDF text: {path}")
    pages_text = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            pages_text.append(text)

    raw = "\n".join(pages_text)

    # Repeated header/footer lines are the main boilerplate in this PDF
    boilerplate = [
        "Horizon Community Health Plan - SBC",
        "Confidential Sample Document",
    ]
    return normalize_text(raw, drop_lines_containing=boilerplate)


# ---------------------------------------------------------------------------
# 2. DOCX extraction (python-docx)
# ---------------------------------------------------------------------------

def extract_docx(path: str) -> str:
    print(f"[2/4] Extracting DOCX text: {path}")
    document = docx.Document(path)

    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # skip fully empty rows (blank fillable fields)
            if any(cells):
                parts.append(" | ".join(cells))

    raw = "\n".join(parts)
    return normalize_text(raw)


# ---------------------------------------------------------------------------
# 3. OCR extraction (pytesseract) for scanned images / scanned PDFs
# ---------------------------------------------------------------------------

def extract_ocr(path: str) -> str:
    print(f"[3/4] Running OCR on scanned document: {path}")

    if path.lower().endswith(".pdf"):
        from pdf2image import convert_from_path
        images = convert_from_path(path)
    else:
        images = [Image.open(path)]

    raw_pages = []
    for i, image in enumerate(images, start=1):
        raw_pages.append(pytesseract.image_to_string(image))

    raw = "\n".join(raw_pages)
    raw = fix_ocr_pipe_i(raw)
    return normalize_text(raw)


# ---------------------------------------------------------------------------
# 4. Web scraping (requests + BeautifulSoup)
# ---------------------------------------------------------------------------

def extract_web(url: str) -> str:
    print(f"[4/4] Scraping public FAQ page: {url}")
    headers = {"User-Agent": "Mozilla/5.0 (educational data-extraction mission)"}
    resp = requests.get(url, headers=headers, timeout=15)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "html.parser")

    # Strip elements that are pure boilerplate/navigation, not content
    for tag in soup(["script", "style", "nav", "footer", "header", "svg", "form"]):
        tag.decompose()

    # Prefer <main> if present, since it usually holds the actual article body
    main = soup.find("main") or soup

    raw = main.get_text(separator="\n")

    nav_boilerplate = [
        "login", "search", "contact us", "about us", "careers",
        "facebook.com", "instagram.com", "twitter.com", "linkedin.com",
        "youtube.com", "tiktok.com", "site map", "language assistance",
        "espanol", "important information", "resources and tools",
        "our sites", "©", "privacy practices", "code of conduct",
    ]
    return normalize_text(raw, drop_lines_containing=nav_boilerplate)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    # 1. PDF -> benefits.txt (policy/benefits document)
    pdf_text = extract_pdf(os.path.join(SOURCE_DIR, "policy.pdf"))
    save_text("benefits.txt", pdf_text)

    # 2. DOCX -> claims_process.txt (claim form / claims process document)
    docx_text = extract_docx(os.path.join(SOURCE_DIR, "claim_form.docx"))
    save_text("claims_process.txt", docx_text)

    # 3. OCR -> enrollment.txt (scanned enrollment form)
    ocr_text = extract_ocr(os.path.join(SOURCE_DIR, "enrollment_scanned.jpg"))
    save_text("enrollment.txt", ocr_text)

    # 4. Web scrape - public BCBS Michigan FAQ page on deductibles/coinsurance/copays
    faq_url = "https://www.bcbsm.com/individuals/help/how-health-insurance-works/deductibles-coinsurance-copays/"
    try:
        faq_text = extract_web(faq_url)
        save_text("provider_faq.txt", faq_text)
    except requests.RequestException as e:
        print(f"  !! Web scrape failed ({e}); network access may be restricted in this environment.")

    print("\nDone. Extracted files are in raw_text/")
